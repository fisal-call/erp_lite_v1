"""
generate_stage_report.py
Generates the ERP Lite Frontend stage status report as a .docx file.

This is an INTERNAL status report (not a client-facing deliverable), so the
formatting is intentionally clean and minimal: A4, Arabic-RTL body, headings,
bullet lists, simple tables. No cover-recipe ceremony.

Arabic font handling: the system has no dedicated Arabic font installed,
so we rely on DejaVu Sans (which has full Arabic coverage) + FreeSerif as
fallback. Word/LibreOffice on the user's machine will substitute a real
Arabic font (Calibri/Cairo/Times New Roman with Arabic support) when opening
the .docx — what we control here is the rtl=True flag and the bidi-visual
paragraph direction.

Usage:
    python3 generate_stage_report.py --stage 1 --title "..." --sections-file sections.json
"""
from __future__ import annotations

import argparse
import json
import sys
from pathlib import Path
from typing import Any

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor


# Font: we set both ascii and eastAsia/cs (complex script). On a machine with
# a real Arabic font, the cs font is what renders Arabic glyphs. We name
# "Arial" and "Calibri" because they ship with Arabic coverage on Windows/
# macOS/Office; if not found, the OS substitutes.
LATIN_FONT = "Calibri"
ARABIC_FONT = "Arial"


def _set_rtl(paragraph) -> None:
    """Force RTL paragraph direction (Arabic)."""
    pPr = paragraph._p.get_or_add_pPr()
    pPr.set(qn("w:bidi"), "1")


def _set_run_rtl_and_font(run, font_name: str = ARABIC_FONT, size_pt: int = 11) -> None:
    run.font.name = LATIN_FONT
    rPr = run._r.get_or_add_rPr()
    # Force complex-script font (Arabic uses cs, not ascii).
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        from docx.oxml import OxmlElement

        rFonts = OxmlElement("w:rFonts")
        rPr.insert(0, rFonts)
    rFonts.set(qn("w:ascii"), LATIN_FONT)
    rFonts.set(qn("w:hAnsi"), LATIN_FONT)
    rFonts.set(qn("w:cs"), font_name)
    rFonts.set(qn("w:eastAsia"), font_name)
    # Mark as RTL run
    rtl = rPr.find(qn("w:rtl"))
    if rtl is None:
        from docx.oxml import OxmlElement

        rtl = OxmlElement("w:rtl")
        rtl.set(qn("w:val"), "1")
        rPr.append(rtl)
    run.font.size = Pt(size_pt)


def add_heading(doc: Document, text: str, level: int = 1) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    _set_rtl(p)
    if level == 1:
        p.style = doc.styles["Heading 1"]
        size = 18
    elif level == 2:
        p.style = doc.styles["Heading 2"]
        size = 15
    else:
        p.style = doc.styles["Heading 3"]
        size = 13
    run = p.add_run(text)
    run.bold = True
    _set_run_rtl_and_font(run, size_pt=size)
    # Heading color: deep navy
    run.font.color.rgb = RGBColor(0x1F, 0x3A, 0x5F)


def add_paragraph(doc: Document, text: str, italic: bool = False) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    _set_rtl(p)
    run = p.add_run(text)
    run.italic = italic
    _set_run_rtl_and_font(run, size_pt=11)


def add_bullet(doc: Document, text: str) -> None:
    p = doc.add_paragraph(style="List Bullet")
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    _set_rtl(p)
    run = p.add_run(text)
    _set_run_rtl_and_font(run, size_pt=11)


def add_code_block(doc: Document, code: str) -> None:
    """Render a code block in monospace, LTR (code is English)."""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    # Do NOT set bidi — keep LTR for code
    run = p.add_run(code)
    run.font.name = "Courier New"
    rPr = run._r.get_or_add_rPr()
    from docx.oxml import OxmlElement

    rFonts = OxmlElement("w:rFonts")
    rFonts.set(qn("w:ascii"), "Courier New")
    rFonts.set(qn("w:hAnsi"), "Courier New")
    rFonts.set(qn("w:cs"), "Courier New")
    existing = rPr.find(qn("w:rFonts"))
    if existing is not None:
        rPr.remove(existing)
    rPr.insert(0, rFonts)
    run.font.size = Pt(9)
    # Light gray background for the code paragraph
    pPr = p._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), "F5F5F5")
    pPr.append(shd)


def add_table(doc: Document, headers: list[str], rows: list[list[str]]) -> None:
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Light Grid Accent 1"
    table.alignment = WD_TABLE_ALIGNMENT.RIGHT
    # Header row
    hdr = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = ""
        p = hdr[i].paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        _set_rtl(p)
        run = p.add_run(h)
        run.bold = True
        _set_run_rtl_and_font(run, size_pt=10)
    # Data rows
    for r, row in enumerate(rows, start=1):
        cells = table.rows[r].cells
        for i, val in enumerate(row):
            cells[i].text = ""
            p = cells[i].paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            _set_rtl(p)
            run = p.add_run(val)
            _set_run_rtl_and_font(run, size_pt=10)


def set_doc_rtl(doc: Document) -> None:
    """Set the entire document to RTL."""
    sectPr = doc.sections[0]._sectPr
    from docx.oxml import OxmlElement

    bidi = OxmlElement("w:bidi")
    sectPr.append(bidi)


def set_page_margins(doc: Document) -> None:
    s = doc.sections[0]
    s.top_margin = Cm(2.5)
    s.bottom_margin = Cm(2.5)
    s.left_margin = Cm(2.5)
    s.right_margin = Cm(2.5)


def build_report(
    output_path: Path,
    title: str,
    subtitle: str,
    sections: list[dict[str, Any]],
) -> None:
    doc = Document()
    set_page_margins(doc)
    set_doc_rtl(doc)

    # Title block
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _set_rtl(p)
    run = p.add_run(title)
    run.bold = True
    _set_run_rtl_and_font(run, size_pt=22)
    run.font.color.rgb = RGBColor(0x1F, 0x3A, 0x5F)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _set_rtl(p)
    run = p.add_run(subtitle)
    _set_run_rtl_and_font(run, size_pt=12)
    run.italic = True
    run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

    doc.add_paragraph()  # spacer

    for section in sections:
        add_heading(doc, section["heading"], level=section.get("level", 1))
        for block in section["blocks"]:
            kind = block["kind"]
            if kind == "paragraph":
                add_paragraph(doc, block["text"], italic=block.get("italic", False))
            elif kind == "bullet":
                add_bullet(doc, block["text"])
            elif kind == "code":
                add_code_block(doc, block["text"])
            elif kind == "table":
                add_table(doc, block["headers"], block["rows"])
            elif kind == "spacer":
                doc.add_paragraph()

    output_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(output_path))
    print(f"✓ Wrote {output_path}")


def main() -> int:
    parser = argparse.ArgumentParser()
    parser.add_argument("--sections-file", required=True, help="JSON file with sections")
    parser.add_argument("--output", required=True, help="Output .docx path")
    parser.add_argument("--title", required=True)
    parser.add_argument("--subtitle", required=True)
    args = parser.parse_args()

    sections = json.loads(Path(args.sections_file).read_text(encoding="utf-8"))
    build_report(Path(args.output), args.title, args.subtitle, sections)
    return 0


if __name__ == "__main__":
    sys.exit(main())
